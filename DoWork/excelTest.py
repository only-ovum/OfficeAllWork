import pandas
import pandas as pd
import docx

#读取文件,用data变量存储读取数据，print打印
# data = pd.read_excel('D:\\文档暂存\\gPTP以太网时间同步测试-已评审，未上传.xlsx')
# 打印某一列的方法
# print(data['测试步骤'])


file = docx.Document('D:\\work_project\\DIDI\\滴滴PBOX软件详细设计 _V1.10.docx')
paragraph = file.paragraphs

for paragraph in file.paragraphs:
    if paragraph.style.name == "Heading 3":
        print(paragraph.text)
